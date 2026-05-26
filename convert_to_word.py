
"""
将Markdown项目报告书转换为Word文档
使用前请先安装依赖：
pip install python-docx markdown
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import markdown
import re

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_heading_font(run):
    """设置标题字体"""
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def parse_markdown_to_docx(md_file, output_file):
    """解析Markdown文件并保存为Word文档"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按行处理
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_content:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # 处理标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            for run in p.runs:
                set_heading_font(run)
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs:
                set_heading_font(run)
                run.font.size = Pt(16)
                run.font.bold = True
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            for run in p.runs:
                set_heading_font(run)
                run.font.size = Pt(14)
                run.font.bold = True
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            for run in p.runs:
                set_heading_font(run)
                run.font.size = Pt(13)
                run.font.bold = True
        elif line.strip() == '---':
            # 分隔线
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.strip() == '':
            # 空行
            doc.add_paragraph()
        elif line.startswith('- ') or line.startswith('* '):
            # 无序列表
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                set_chinese_font(run)
        elif re.match(r'^\d+\.\s', line):
            # 有序列表
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                set_chinese_font(run)
        elif line.startswith('|') and '|' in line[1:]:
            # 简单表格处理（这里简化处理，只添加为段落）
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_chinese_font(run)
        else:
            # 普通段落
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_chinese_font(run)
    
    doc.save(output_file)
    print(f"Word文档已生成：{output_file}")

if __name__ == '__main__':
    try:
        parse_markdown_to_docx('项目报告书.md', '项目报告书.docx')
        print("\n转换成功！您可以使用Word打开项目报告书.docx")
        print("建议：打开后可以根据需要调整格式和插入图片")
    except ImportError as e:
        print(f"缺少依赖库，请先安装：pip install python-docx")
        print("错误信息：{e}")
    except Exception as e:
        print(f"转换过程出错：{e}")

